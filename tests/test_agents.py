import io

import docx
import pytest
from fastapi.testclient import TestClient

from app.config import get_settings
from app.main import app


@pytest.fixture(scope="module")
def settings():
    return get_settings()


@pytest.fixture(scope="module")
def test_client(settings):
    with TestClient(app) as client:
        yield client


def test_create_agent(test_client):
    data = {"name": "TestAgent"}
    response = test_client.post("/agents", data=data)
    assert response.status_code == 201
    json_data = response.json()
    # Expecting a response of the form {"id": "<agent_id>"}
    assert "id" in json_data
    id = json_data["id"]
    test_client.delete(f"/agents/{id}")


def test_get_all_agents(test_client):
    response = test_client.get("/agents")
    assert response.status_code == 200
    agents = response.json()
    assert isinstance(agents, list)


def test_create_get_and_delete_agent(test_client):
    data = {"name": "TempAgent"}
    response = test_client.post("/agents", data=data)
    assert response.status_code == 201
    agent_id = response.json()["id"]

    response = test_client.get(f"/agents/{agent_id}")
    assert response.status_code == 200
    agent = response.json()
    assert agent["name"] == "TempAgent"

    response = test_client.delete(f"/agents/{agent_id}")
    assert response.status_code == 204

    response = test_client.get(f"/agents/{agent_id}")
    assert response.status_code == 404


def test_update_agent_websites(test_client):
    data = {"name": "WebAgent"}
    response = test_client.post("/agents", data=data)
    agent_id = response.json()["id"]

    websites = ["https://fastapi.tiangolo.com/tutorial/testing/#extended-testing-file"]
    response = test_client.put(f"/agents/{agent_id}/websites", json=websites)
    assert response.status_code == 204

    response = test_client.get(f"/agents/{agent_id}")
    assert response.status_code == 200

    agent = response.json()
    assert "files" in agent
    assert len(agent["files"]) > 0
    test_client.delete(f"/agents/{agent_id}")


def test_update_agent_files(test_client):
    data = {"name": "FileAgent"}
    response = test_client.post("/agents", data=data)
    agent_id = response.json()["id"]

    files = []
    for filename in ["Resume.pdf", "HowToReadAPaper.png"]:
        with open(f"tests/docs/{filename}", "rb") as f:
            file_content = f.read()
            content_type = (
                "application/pdf" if filename.endswith(".pdf") else "image/png"
            )
            files.append(("files", (filename, io.BytesIO(file_content), content_type)))

    response = test_client.put(f"/agents/{agent_id}/files", files=files)
    assert response.status_code == 204

    response = test_client.get(f"/agents/{agent_id}")
    assert response.status_code == 200
    agent = response.json()
    assert "files" in agent
    assert any("Resume.pdf" in f.get("name", "") for f in agent["files"])


def test_send_message(test_client):
    data = {"name": "QueryAgent"}
    response = test_client.post("/agents", data=data)
    agent_id = response.json()["id"]

    with open("tests/docs/Resume.pdf", "rb") as f:
        file_content = f.read()
    files = [("files", ("Resume.pdf", io.BytesIO(file_content), "application/pdf"))]
    response = test_client.put(f"/agents/{agent_id}/files", files=files)
    assert response.status_code == 204

    payload = {"message": "Who is Aaron Tua and what is he known for?"}
    response = test_client.post(f"/agents/{agent_id}/queries", json=payload)
    assert response.status_code == 201
    result = response.json()
    assert isinstance(result, dict)
    assert "messages" in result
    test_client.delete(f"/agents/{agent_id}")


def test_exceed_token_limit(test_client):
    data = {"name": "TokenLimitAgent"}
    response = test_client.post("/agents", data=data)
    agent_id = response.json()["id"]

    document = docx.Document()
    document.add_paragraph("Hello world!", style="Heading 1")
    document.add_paragraph("This is a test document." * 100000, style="Normal")
    file_content = io.BytesIO()
    document.save(file_content)
    file_content.seek(0)
    files = [
        (
            "files",
            (
                "overloaded.docx",
                file_content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )
    ]
    response = test_client.put(f"/agents/{agent_id}/files", files=files)
    assert response.status_code == 422
    test_client.delete(f"/agents/{agent_id}")
